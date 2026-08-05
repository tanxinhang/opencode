"""Append the recent multistatic receiver Gate results to all model docs.

The script appends a clearly labelled Chinese appendix to the revised
argument/system model and to the System_Model_revised document.  It removes an
existing "附录 A" first, so repeated runs stay idempotent.
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


CHINESE_SOURCE = Path("UAV_OTFS_ISAC_论证与系统模型_revised_final.docx")
CHINESE_OUTPUTS = (
    Path("UAV_OTFS_ISAC_论证与系统模型_revised_final.docx"),
    Path("UAV_OTFS_ISAC_论证与系统模型_revised_final_G0C.docx"),
)
ENGLISH_SOURCE = Path("UAV_OTFS_ISAC_System_Model_revised.docx")


def set_font(run, size=10.5, bold=False, color="000000"):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def append_to_body_end(doc, paragraph):
    body = doc._body._element
    section_breaks = body.findall(qn("w:sectPr"))
    if section_breaks:
        section_breaks[-1].addprevious(paragraph._p)
    else:
        body.append(paragraph._p)


def body(doc, text):
    p = doc.add_paragraph()
    set_font(p.add_run(text))
    append_to_body_end(doc, p)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_font(p.add_run(text))
    append_to_body_end(doc, p)
    return p


def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    append_to_body_end(doc, p)
    return p


def remove_existing_appendix(doc) -> None:
    index = next((
        index for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.strip().startswith(("附录 A", "附录 B", "附录 C"))
    ), None)
    if index is None:
        return
    for paragraph in doc.paragraphs[index:]:
        paragraph._p.getparent().remove(paragraph._p)


def add_chinese_appendix(doc) -> None:
    heading(doc, "附录 A：多站 OTFS 接收机 Gate 更新", 1)
    body(doc, "本附录记录系统模型定稿后新增的接收机级 Gate 与性能审计结果，内容与仓库 README 的 Gate G0-C 部分一致。")
    heading(doc, "A.1 Gate G0-C：玩具 MF/CFAR 前端接入关联后端", 2)
    bullet(doc, "八单元 ULA 接收多 UAV 单位能量 QPSK 身份签名，做可分离角度-时延-多普勒匹配滤波；")
    bullet(doc, "采用经验 max-map 阈值控制帧级虚警，并用 isotonic 回归把峰值分数校准为路径存在概率；")
    bullet(doc, "逐视角假目标/假额外峰概率在保留的单目标成分上估计，碰撞支持阈值允许外部覆盖。")
    heading(doc, "A.2 每路径 Fisher 型协方差", 2)
    bullet(doc, "由匹配滤波能量峰局部曲率导出角度、距离、多普勒标准差；")
    bullet(doc, "后端 GLS 位置重估和 Huber 多普勒重估使用逐路径精度；")
    bullet(doc, "配对 50 帧结果：检测决策不变，GOSPA 明显下降，位置误差约从 1.1-1.4 m 降至 0.6-0.8 m。")
    heading(doc, "A.3 多帧非相干积累与旁瓣参考 CFAR", 2)
    bullet(doc, "四帧能量平均实现非相干积累，弱路径检测概率提高；")
    bullet(doc, "阈值在保留单目标帧上挖空真实峰值邻域后校准，压在确定性旁瓣地板之上；")
    bullet(doc, "30 帧结果：N=1 与 N=2 场景精确恢复均达到 96.7%，H1 假峰降至 0-0.1/帧。")
    heading(doc, "A.4 等资源与开放边界", 2)
    bullet(doc, "等总导频能量审计表明：把每帧幅度减半再做四帧积累不会保留增益，收益来自帧/能量预算；")
    bullet(doc, "独立瑞利衰落首轮等平均能量审计未恢复增益，时间分集暂不作为已证明的公平增益；")
    bullet(doc, "同一 angle-DD 单元内碰撞分解、强 FWER、带宽/帧预算/通信速率对账仍为开放问题。")
    heading(doc, "附录 B：G1 路线与表述修正", 1)
    body(doc, "G1 系列 Gate 用于把 G0-C 的物理证据接入论文主线的可靠性与相关性校准选择，并修正此前对 Top-K 的笼统表述。")
    heading(doc, "B.1 G1-A 证据矩校准", 2)
    bullet(doc, "从 G0-C 前端按 H0/H1 导出每 UAV 证据 z_iq，估计 (mu_h, Sigma_h)，要求收缩后协方差正定；")
    bullet(doc, "验证预测单链路 deflection 与固定 P_FA 下实际 P_D 排序一致，建议 Spearman 相关不低于 0.6；")
    bullet(doc, "正式统计（5000 训练/5000 测试几何）：deflection 作为预测分时 Spearman 0.588，低于 0.6；改用精确 P_D 增益预测分后正式 10k Spearman 0.996（CI [0.98,1.00]），logit/相对漏检缺口 0.994，P_D 增益选择器正式通过 G1-A。")
    heading(doc, "B.2 G1-B 量化与报告信道闭环", 2)
    bullet(doc, "对量化比特、BER/转移矩阵、可检测擦除和实际收到集合重新融合；")
    bullet(doc, "Monte Carlo 矩与精确公式的均值最大相对误差 4.08%、对角与主交叉协方差最大 8.51%，满足 5%/10% 目标；")
    bullet(doc, "扫描范围：量化比特 1-4、BER 0.01/0.08、擦除 0.9/0.7、相关性 0/0.5/0.9。")
    heading(doc, "B.3 G1-C 条件重排序价值", 2)
    bullet(doc, "当前方法定义为 Conditional-Deflection Greedy，候选得分随已选集合变化；")
    bullet(doc, "不声称“优于 Top-K”，只验证条件重排序相对 Static ID Top-K 的可测量增益；")
    bullet(doc, "协方差对角、成本与成功概率相同时，方法应退化为静态单链路 deflection Top-K；")
    bullet(doc, "当前 smoke 通过退化一致性，并在高相关 vs 低相关场景中选择低相关报告且获得更高 P_D。")
    heading(doc, "B.4 G1-D 贪婪近似 vs Oracle", 2)
    bullet(doc, "8 组配置 smoke：开环 pi*Delta-D/b 与精确边际期望 deflection 增益的 Spearman 为 0.90；")
    bullet(doc, "一阶、精确和 SAA 贪婪与穷举 Oracle 的选择一致率均为 50%，平均期望 deflection 缺口 0.161；")
    bullet(doc, "结论：一阶评分排序良好，但预算与集合交互仍会带来 Oracle 差距，不能直接视为无约束最优。")
    heading(doc, "B.5 G2 系统级预算扫描", 2)
    bullet(doc, "N=8/12、Q=3/5、B_max=20/40、20 种子的公平全局预算扫描：Proposed 平均 P_D 0.898（最差 0.814），Sensing-SNR Top-K 0.898，Independent-Deflection Top-K 0.897，Communication Top-K 0.773，All-scheduled 0.935；")
    bullet(doc, "此前“Proposed 落后 3.6 pp”来自不公平的逐目标预算基线，已修正；")
    bullet(doc, "精确 P_D 增益贪心平均 P_D 0.900，为贪心变体中最强；J 散度替代目标在异方差矩下与 P_D 不一致，已作为负结果记录；")
    bullet(doc, "强相关模型（最高 SNR 报告对 rho=0.85、20 种子）下，条件贪心平均 P_D 0.870 高于 Independent-Deflection Top-K 0.855，在 83.1% 配置中胜出；精确 P_D 贪心 0.880 为最强变体，支持将选择器升级为 P_D 增益贪心；")
    bullet(doc, "多 rho 扫描（0/0.3/0.5/0.7/0.85）显示条件贪心在 rho>=0.3 多数单元有正的配对差 CI（如 rho=0.5、B=20 时 +0.0199，CI [0.013,0.027]），rho=0.85、B=20 时 CI 跨零；精确 P_D 贪心在每个 rho 均最强。")
    heading(doc, "B.6 创新定位", 2)
    bullet(doc, "创新点是集成场景与端到端验证链，而不是新的选择算法家族；")
    bullet(doc, "条件边际 deflection 贪心是对既有 deflection 最优线性融合与贪心子集选择的适配，不声称新算法，也不声称普遍优于 Top-K；")
    bullet(doc, "若投稿要求算法创新，需要升级为例如 logit-P_D 增益贪心并给出形式化选择性质，仅场景创新不足以支撑算法层面的新贡献。")
    heading(doc, "附录 C：论文写作骨架", 1)
    body(doc, "对应仓库 PAPER_OUTLINE.md，用于把现有 Gate 结果整理为可投稿结构。")
    bullet(doc, "标题方向：通信损伤相关软证据下的 UAV-OTFS-ISAC 选择性融合；")
    bullet(doc, "创新定位：新场景与端到端验证链，算法为既有条件重排序的适配；")
    bullet(doc, "核心证据：强相关模型下条件贪心在 77.5% 配置中超过静态 Independent-Deflection Top-K；")
    bullet(doc, "投稿前必须补：G1-A 万帧统计、G2 10-20 种子胜率 CI、等帧/等能量/等帧预算资源对账。")


def main() -> None:
    for source, outputs in (
        (CHINESE_SOURCE, CHINESE_OUTPUTS),
        (ENGLISH_SOURCE, (ENGLISH_SOURCE,)),
    ):
        if not source.exists():
            raise FileNotFoundError(source)
        doc = Document(source)
        remove_existing_appendix(doc)
        add_chinese_appendix(doc)
        for output in outputs:
            output.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output)
            print(output)


if __name__ == "__main__":
    main()
