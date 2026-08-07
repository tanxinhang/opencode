"""Structural audit for paper/submission.docx."""

from __future__ import annotations

from pathlib import Path
import sys

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "paper" / "submission.docx"


def main() -> None:
    if not DOCX.exists():
        print(f"missing {DOCX}; run scripts/md_to_docx.py first")
        sys.exit(1)
    doc = Document(DOCX)
    checks = [
        ("title present",
         any(p.style.name == "Title" for p in doc.paragraphs)),
        ("at least 20 headings",
         sum(1 for p in doc.paragraphs
             if p.style.name.startswith("Heading")) >= 20),
        ("six tables", len(doc.tables) == 6),
        ("three inline images", len(doc.inline_shapes) == 3),
        ("22 references",
         any(p.text.strip().startswith("[22]") for p in doc.paragraphs)),
    ]
    section = doc.sections[0]
    checks.append(
        ("US Letter margins 1 in",
         abs(section.page_width.inches - 8.5) < 1e-6
         and abs(section.page_height.inches - 11.0) < 1e-6
         and abs(section.left_margin.inches - 1.0) < 1e-6
         and abs(section.top_margin.inches - 1.0) < 1e-6)
    )
    failed = False
    for name, passed in checks:
        print(f"[{'PASS' if passed else 'FAIL'}] {name}")
        failed = failed or not passed
    if failed:
        sys.exit(1)
    print("submission.docx structural audit passed.")


if __name__ == "__main__":
    main()
