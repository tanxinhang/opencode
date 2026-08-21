from docx import Document
import sys, json

for path in sys.argv[1:]:
    d = Document(path)
    print("\nFILE", path)
    print("PARAGRAPHS", len(d.paragraphs), "TABLES", len(d.tables))
    for i, p in enumerate(d.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith("Heading") or p.style.name == "Equation":
            print(f"{i:03d}\t{p.style.name}\t{text}")
    print("TABLE TEXT")
    for ti, t in enumerate(d.tables):
        rows = [[c.text.replace('\n',' / ') for c in r.cells] for r in t.rows]
        print(ti, json.dumps(rows[:4], ensure_ascii=False))
