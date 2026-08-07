"""Convert paper/submission.md into paper/submission.docx.

The converter is tailored to the repository manuscript: Markdown headings,
paragraphs, bullet/numbered lists, pipe tables, images, bold labels, and
LaTeX-style math are mapped to Word paragraphs, tables, and pictures.
Mathematical formulas are kept as readable plain text because a Word
equation renderer is not available in this environment.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "paper" / "submission.md"
OUT = ROOT / "paper" / "submission.docx"
FIGURES = ROOT / "paper_figures"

HEADING_COLORS = {
    0: "111827",
    1: "2E74B5",
    2: "2E74B5",
    3: "1F4D78",
}


def add_inline_runs(paragraph, text: str, base_bold: bool = False):
    """Add runs with **bold** and `code` inline markup."""
    pattern = re.compile(r"(\*\*.*?\*\*|`[^`]*`)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold


def set_style_geometry(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for level, color in HEADING_COLORS.items():
        style = doc.styles["Title"] if level == 0 else doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        if level == 0:
            style.font.size = Pt(17)
            style.paragraph_format.space_after = Pt(8)
        elif level == 1:
            style.font.size = Pt(16)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        elif level == 2:
            style.font.size = Pt(13)
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(4)
        else:
            style.font.size = Pt(12)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    clean = [
        [re.sub(r"\*\*", "", cell.strip()) for cell in row]
        for row in rows
    ]
    if not clean:
        return
    table = doc.add_table(rows=len(clean), cols=len(clean[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = Inches(6.5)
    col_width = width / len(clean[0])
    for row_index, row in enumerate(clean):
        for col_index, text in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.width = col_width
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(text)
            run.font.size = Pt(9)
            if row_index == 0:
                run.bold = True
                shading = cell._tc.get_or_add_tcPr().makeelement(
                    qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "E8EEF5"}
                )
                cell._tc.get_or_add_tcPr().append(shading)


def add_image(doc: Document, path: str) -> None:
    image_path = FIGURES / Path(path).name
    if image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.0))


def main() -> None:
    doc = Document()
    set_style_geometry(doc)

    lines = SRC.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            paragraph = doc.add_paragraph(style="Title")
            add_inline_runs(paragraph, stripped[2:])
        elif stripped.startswith("```"):
            code: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != "```":
                code.append(lines[index])
                index += 1
            paragraph = doc.add_paragraph()
            for code_index, code_line in enumerate(code):
                run = paragraph.add_run(code_line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                if code_index < len(code) - 1:
                    run.add_break()
            index += 1
        elif stripped.startswith("**Keywords:**"):
            parts = [stripped[len("**Keywords:**"):].strip()]
            while (
                index + 1 < len(lines)
                and lines[index + 1].strip()
            ):
                index += 1
                parts.append(lines[index].strip())
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("Keywords: ")
            run.bold = True
            add_inline_runs(paragraph, " ".join(parts))
            index += 1
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline_runs(paragraph, stripped[3:])
        elif stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline_runs(paragraph, stripped[4:])
        elif stripped.startswith("![") and "](paper_figures/" in stripped:
            path = stripped.split("](", 1)[1].rstrip(")")
            add_image(doc, path)
        elif stripped.startswith("**Table ") or stripped.startswith("**Figure "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            add_inline_runs(paragraph, stripped, base_bold=True)
        elif stripped.startswith("|"):
            table_rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [
                    cell.strip()
                    for cell in lines[index].strip().strip("|").split("|")
                ]
                if not all(re.fullmatch(r":?-{2,}:?", cell) for cell in cells):
                    table_rows.append(cells)
                index += 1
            add_table(doc, table_rows)
            continue
        elif stripped.startswith("$$"):
            math_lines = []
            if stripped == "$$":
                index += 1
                while index < len(lines) and lines[index].strip() != "$$":
                    math_lines.append(lines[index].strip())
                    index += 1
            else:
                math_lines.append(stripped[2:])
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(" ".join(math_lines))
            run.font.name = "Cambria Math"
            run.font.size = Pt(10)
        elif stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_runs(paragraph, re.sub(r"^\d+\.\s", "", stripped))
        else:
            paragraph_lines = [stripped]
            while index + 1 < len(lines):
                next_line = lines[index + 1].rstrip()
                next_stripped = next_line.strip()
                if not next_stripped:
                    break
                if (
                    next_stripped.startswith("#")
                    or next_stripped.startswith("|")
                    or next_stripped.startswith("![")
                    or next_stripped.startswith("**Table ")
                    or next_stripped.startswith("**Figure ")
                    or next_stripped.startswith("**Keywords:")
                    or next_stripped.startswith("**Theorem ")
                    or next_stripped.startswith("**Lemma ")
                    or next_stripped.startswith("$$")
                    or next_stripped.startswith("```")
                    or next_stripped.startswith("- ")
                    or re.match(r"^\d+\.\s", next_stripped)
                ):
                    break
                paragraph_lines.append(next_stripped)
                index += 1
            paragraph = doc.add_paragraph()
            add_inline_runs(paragraph, " ".join(paragraph_lines))
        index += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
